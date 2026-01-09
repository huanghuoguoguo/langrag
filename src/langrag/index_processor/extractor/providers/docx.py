"""DOCX file parser"""

from __future__ import annotations

from pathlib import Path

from loguru import logger

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. DOCX parsing will not be available.")

from langrag.entities.document import Document

from ..base import BaseParser


class DocxParser(BaseParser):
    """DOCX file parser

    Uses python-docx to extract text content from Word documents.

    Args:
        include_tables: Whether to include table content
        include_headers: Whether to include header content
        include_footers: Whether to include footer content

    Usage example:
        >>> parser = DocxParser()
        >>> docs = parser.parse("document.docx")
    """

    def __init__(
        self,
        include_tables: bool = True,
        include_headers: bool = False,
        include_footers: bool = False,
    ):
        """Initialize DOCX parser

        Args:
            include_tables: Whether to include tables
            include_headers: Whether to include headers
            include_footers: Whether to include footers
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX parsing. Install it with: pip install python-docx"
            )

        self.include_tables = include_tables
        self.include_headers = include_headers
        self.include_footers = include_footers

    def parse(self, file_path: str | Path, **_kwargs) -> list[Document]:
        """Parse DOCX file

        Args:
            file_path: DOCX file path
            **kwargs: Additional parameters

        Returns:
            A list containing a single Document

        Raises:
            FileNotFoundError: File does not exist
            ValueError: Not a valid DOCX file
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        if not path.is_file():
            raise ValueError(f"Not a file: {path}")

        logger.info(f"Parsing DOCX file: {path}")

        try:
            doc = DocxDocument(str(path))
            text_content = []

            # Extract headers
            if self.include_headers:
                for section in doc.sections:
                    header = section.header
                    for paragraph in header.paragraphs:
                        if paragraph.text.strip():
                            text_content.append(f"[Header] {paragraph.text}")

            # Extract paragraphs
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
                    paragraph_count += 1

            # Extract tables
            table_count = 0
            if self.include_tables:
                for table in doc.tables:
                    table_text = self._extract_table(table)
                    if table_text:
                        text_content.append(table_text)
                        table_count += 1

            # Extract footers
            if self.include_footers:
                for section in doc.sections:
                    footer = section.footer
                    for paragraph in footer.paragraphs:
                        if paragraph.text.strip():
                            text_content.append(f"[Footer] {paragraph.text}")

            content = "\n".join(text_content)

            if not content.strip():
                logger.warning(f"No text content extracted from {path}")

            doc_obj = Document(
                page_content=content,
                metadata={
                    "source": str(path.absolute()),
                    "filename": path.name,
                    "extension": path.suffix,
                    "paragraphs": paragraph_count,
                    "tables": table_count,
                    "parser": "DocxParser",
                },
            )

            logger.info(
                f"Parsed DOCX: {paragraph_count} paragraphs, "
                f"{table_count} tables, {len(content)} characters"
            )

            return [doc_obj]

        except Exception as e:
            logger.error(f"Failed to parse DOCX {path}: {e}")
            raise ValueError(f"Invalid DOCX file: {e}") from e

    def _extract_table(self, table) -> str:
        """Extract table content to Markdown format

        Args:
            table: python-docx Table object

        Returns:
            Table string in Markdown format
        """
        lines = []

        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))

            # Add separator after the first row
            if i == 0:
                lines.append(" | ".join(["---"] * len(cells)))

        return "\n".join(lines) if lines else ""
